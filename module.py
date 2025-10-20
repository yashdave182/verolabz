import json
import re
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Any, Optional, Tuple
from enum import Enum
from collections import defaultdict

# Import Word document libraries
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# DOCUMENT TWEAK - FORMAT TEMPLATE SYSTEM
# ============================================================================

# ============================================================================
# CORE DATA STRUCTURES
# ============================================================================

class FormatType(Enum):
    BOLD = "bold"
    ITALIC = "italic"
    UNDERLINE = "underline"
    FONT_SIZE = "font_size"
    FONT_FAMILY = "font_family"
    COLOR = "color"
    BACKGROUND = "background"
    ALIGNMENT = "alignment"

class BlockType(Enum):
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    LIST_ITEM = "list_item"
    TABLE_CELL = "table_cell"
    CODE = "code"
    QUOTE = "quote"

@dataclass
class TextFormat:
    """Character-level formatting"""
    start: int
    end: int
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_size: int = 12
    font_family: str = "Arial"
    color: str = "#000000"
    
    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(**data)

@dataclass
class BlockFormat:
    """Block-level formatting (paragraph, heading, etc.)"""
    type: BlockType
    alignment: str = "left"  # left, center, right, justify
    margin_top: int = 0
    margin_bottom: int = 0
    margin_left: int = 0
    line_height: float = 1.0
    indent_level: int = 0
    list_style: Optional[str] = None  # bullet, number, none
    heading_level: Optional[int] = None  # 1-6 for headings
    
    def to_dict(self):
        data = asdict(self)
        data['type'] = self.type.value
        return data
    
    @classmethod
    def from_dict(cls, data):
        data['type'] = BlockType(data['type'])
        return cls(**data)

@dataclass
class PositionInfo:
    """Position information for layout preservation"""
    x: float = 0.0
    y: float = 0.0
    width: float = 0.0
    height: float = 0.0
    page_number: int = 1
    
    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(**data)

@dataclass
class RichFormatBlock:
    """A block of text with rich formatting and layout information"""
    text: str
    block_format: BlockFormat
    char_formats: List[TextFormat] = field(default_factory=list)
    position: PositionInfo = field(default_factory=PositionInfo)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self):
        return {
            'text': self.text,
            'block_format': self.block_format.to_dict(),
            'char_formats': [cf.to_dict() for cf in self.char_formats],
            'position': self.position.to_dict(),
            'metadata': self.metadata
        }
    
    @classmethod
    def from_dict(cls, data):
        return cls(
            text=data['text'],
            block_format=BlockFormat.from_dict(data['block_format']),
            char_formats=[TextFormat.from_dict(cf) for cf in data['char_formats']],
            position=PositionInfo.from_dict(data.get('position', {})),
            metadata=data.get('metadata', {})
        )

@dataclass
class RichFormatTemplate:
    """Complete format template with rich layout information"""
    blocks: List[RichFormatBlock] = field(default_factory=list)
    layout: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self):
        return {
            'blocks': [block.to_dict() for block in self.blocks],
            'layout': self.layout,
            'metadata': self.metadata
        }
    
    @classmethod
    def from_dict(cls, data):
        return cls(
            blocks=[RichFormatBlock.from_dict(b) for b in data['blocks']],
            layout=data.get('layout', {}),
            metadata=data.get('metadata', {})
        )
    
    def save(self, filepath: str):
        """Save template to JSON file"""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.to_dict(), f, indent=2, ensure_ascii=False)
    
    @classmethod
    def load(cls, filepath: str):
        """Load template from JSON file"""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return cls.from_dict(data)

# ============================================================================
# FORMAT EXTRACTOR (From plain text and OCR data)
# ============================================================================

class TextFormatExtractor:
    """Extract formatting from plain text and OCR data"""
    
    def extract_from_text_patterns(self, text: str, ocr_data: Optional[Dict] = None) -> RichFormatTemplate:
        """
        Extract format by analyzing text patterns and OCR data if available
        
        Args:
            text: The extracted text
            ocr_data: Optional OCR data with layout information
        """
        template = RichFormatTemplate()
        
        # If we have OCR data with line metadata, use it for better format extraction
        if ocr_data and 'structured_data' in ocr_data and ocr_data['structured_data']:
            return self._extract_from_line_metadata(text, ocr_data['structured_data'])
        
        # Fallback to pattern-based extraction
        lines = text.split('\n')
        
        for line_num, line in enumerate(lines):
            if not line.strip():
                # Preserve empty lines
                template.blocks.append(RichFormatBlock(
                    text="",
                    block_format=BlockFormat(
                        type=BlockType.PARAGRAPH,
                        margin_bottom=10
                    ),
                    position=PositionInfo(y=line_num * 20)  # Approximate positioning
                ))
                continue
            
            block = self._analyze_line_format(line, line_num)
            # Add position information
            block.position = PositionInfo(y=line_num * 20)
            template.blocks.append(block)
        
        return template
    
    def _extract_from_line_metadata(self, text: str, structured_data: Dict) -> RichFormatTemplate:
        """Extract format from line metadata provided by Unstract"""
        template = RichFormatTemplate()
        
        # Store any available layout information
        if 'line_metadata' in structured_data:
            template.metadata['line_metadata'] = structured_data['line_metadata']
        
        if 'highlights' in structured_data:
            template.metadata['highlights'] = structured_data['highlights']
        
        # Try to parse line metadata into rich format blocks
        if 'line_metadata' in structured_data:
            line_metadata = structured_data['line_metadata']
            
            # Process each line with its metadata
            for i, line_info in enumerate(line_metadata.get('lines', [])):
                try:
                    block = self._create_rich_block_from_line_metadata(line_info, i)
                    template.blocks.append(block)
                except Exception as e:
                    print(f"[Format] Warning: Could not parse line metadata {i}: {str(e)}")
                    # Continue with next line
                    continue
        else:
            # Fallback to text-based extraction if no line metadata
            lines = text.split('\n')
            for line_num, line in enumerate(lines):
                if line.strip():
                    block = self._analyze_line_format(line, line_num)
                    block.position = PositionInfo(y=line_num * 20)
                    template.blocks.append(block)
        
        return template
    
    def _create_rich_block_from_line_metadata(self, line_info: Dict, index: int) -> RichFormatBlock:
        """Create a RichFormatBlock from line metadata"""
        # Extract text
        text = line_info.get('text', '')
        
        # Extract position information
        position = PositionInfo()
        if 'bbox' in line_info:
            bbox = line_info['bbox']
            position.x = bbox.get('x', 0)
            position.y = bbox.get('y', 0)
            position.width = bbox.get('width', 0)
            position.height = bbox.get('height', 0)
            position.page_number = bbox.get('page', 1)
        
        # Extract line number if available
        line_number = line_info.get('line_number', index)
        
        # Extract formatting information
        block_format = BlockFormat(type=BlockType.PARAGRAPH)
        char_formats = []
        
        # Determine block type based on content and formatting
        if self._is_heading(text):
            level = self._get_heading_level(text)
            block_format.type = BlockType.HEADING
            block_format.heading_level = level
            block_format.margin_top = 15
            block_format.margin_bottom = 10
        elif self._is_list_item(text):
            block_format.type = BlockType.LIST_ITEM
            block_format.margin_left = 20
            block_format.list_style = "bullet"
        
        # Extract font information if available
        if 'font' in line_info:
            font_info = line_info['font']
            font_family = font_info.get('family', 'Arial')
            font_size = font_info.get('size', 12)
            is_bold = font_info.get('bold', False)
            is_italic = font_info.get('italic', False)
            
            # Add character formatting
            char_formats.append(TextFormat(
                start=0,
                end=len(text),
                font_family=font_family,
                font_size=font_size,
                bold=is_bold,
                italic=is_italic
            ))
        
        return RichFormatBlock(
            text=text,
            block_format=block_format,
            char_formats=char_formats,
            position=position,
            metadata={
                'line_number': line_number,
                'source': 'line_metadata'
            }
        )
    
    def _analyze_line_format(self, line: str, line_num: int = 0) -> RichFormatBlock:
        """Analyze a line and guess its formatting"""
        
        # Check if it's a heading
        if self._is_heading(line):
            level = self._get_heading_level(line)
            return RichFormatBlock(
                text=line.strip(),
                block_format=BlockFormat(
                    type=BlockType.HEADING,
                    heading_level=level,
                    margin_top=15,
                    margin_bottom=10
                ),
                char_formats=[TextFormat(
                    start=0,
                    end=len(line.strip()),
                    bold=True,
                    font_size=24 - (level * 2)
                )]
            )
        
        # Check if it's a list item
        if self._is_list_item(line):
            return RichFormatBlock(
                text=line.strip(),
                block_format=BlockFormat(
                    type=BlockType.LIST_ITEM,
                    margin_left=20,
                    list_style="bullet"
                )
            )
        
        # Check for special markers
        if "<<<PAGE_BREAK>>>" in line:
            # Handle page breaks
            return RichFormatBlock(
                text=line.strip(),
                block_format=BlockFormat(
                    type=BlockType.PARAGRAPH,
                    margin_bottom=20
                )
            )
        
        # Default: paragraph
        return RichFormatBlock(
            text=line.strip(),
            block_format=BlockFormat(
                type=BlockType.PARAGRAPH,
                margin_bottom=10
            )
        )
    
    def _is_heading(self, line: str) -> bool:
        """Detect if line is a heading"""
        line = line.strip()
        
        # All caps
        if line.isupper() and len(line.split()) <= 10:
            return True
        
        # Title case and short
        if line and line[0].isupper() and len(line.split()) <= 8:
            words = line.split()
            if sum(1 for w in words if w[0].isupper()) >= len(words) * 0.7:
                return True
        
        # Ends without punctuation
        if line and line[-1] not in '.!?,:;':
            if len(line.split()) <= 10:
                return True
        
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """Determine heading level (1-6)"""
        if line.isupper():
            return 1
        if len(line.split()) <= 4:
            return 2
        return 3
    
    def _is_list_item(self, line: str) -> bool:
        """Detect if line is a list item"""
        line = line.strip()
        # Bullet points
        if re.match(r'^[-•●○▪▫■□*+]\s', line):
            return True
        # Numbered lists
        if re.match(r'^\d+[\.)]\s', line):
            return True
        return False

# ============================================================================
# FORMAT APPLIER (Apply format to new text)
# ============================================================================

class FormatApplier:
    """Apply formatting template to new text from LLM"""
    
    def apply_format(self, new_text: str, template: RichFormatTemplate) -> RichFormatTemplate:
        """
        Apply the formatting from template to new text
        Maps new text to old formatting structure
        """
        new_template = RichFormatTemplate(
            layout=template.layout.copy(),
            metadata=template.metadata.copy()
        )
        
        # Split new text into blocks (paragraphs)
        new_blocks = self._split_into_blocks(new_text)
        
        # Map new blocks to old format blocks
        for i, new_block_text in enumerate(new_blocks):
            if i < len(template.blocks):
                # Use existing format
                old_block = template.blocks[i]
                new_block = self._apply_block_format(new_block_text, old_block)
            else:
                # Create new block with default format from last block
                if template.blocks:
                    last_format = template.blocks[-1].block_format
                else:
                    last_format = BlockFormat(type=BlockType.PARAGRAPH)
                
                new_block = RichFormatBlock(
                    text=new_block_text,
                    block_format=last_format
                )
            
            new_template.blocks.append(new_block)
        
        return new_template
    
    def apply_format_semantic(self, new_text: str, template: RichFormatTemplate) -> RichFormatTemplate:
        """
        Apply format using semantic matching - better for when LLM changes content structure
        """
        new_template = RichFormatTemplate(
            layout=template.layout.copy(),
            metadata=template.metadata.copy()
        )
        
        # Parse new text into semantic blocks
        new_blocks = self._parse_semantic_blocks(new_text)
        
        # Match new blocks to template using semantic analysis
        matched_blocks = self._match_semantic_blocks(new_blocks, template.blocks)
        
        # Apply formatting
        for new_block, template_block in matched_blocks:
            formatted_block = self._apply_block_format(new_block.text, template_block)
            new_template.blocks.append(formatted_block)
        
        return new_template
    
    def _split_into_blocks(self, text: str) -> List[str]:
        """Split text into logical blocks"""
        # Split by double newlines or single newlines
        blocks = []
        current_block = []
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                if current_block:
                    blocks.append(' '.join(current_block))
                    current_block = []
            else:
                current_block.append(line)
        
        if current_block:
            blocks.append(' '.join(current_block))
        
        return blocks
    
    def _parse_semantic_blocks(self, text: str) -> List[RichFormatBlock]:
        """Parse text into semantic blocks with guessed formatting"""
        lines = text.split('\n')
        blocks = []
        
        extractor = TextFormatExtractor()
        for line_num, line in enumerate(lines):
            if line.strip():
                block = extractor._analyze_line_format(line, line_num)
                block.position = PositionInfo(y=line_num * 20)
                blocks.append(block)
        
        return blocks
    
    def _match_semantic_blocks(self, new_blocks: List[RichFormatBlock], template_blocks: List[RichFormatBlock]) -> List[Tuple[RichFormatBlock, RichFormatBlock]]:
        """Match new blocks to template blocks using semantic analysis"""
        matches = []
        
        # Simple matching: preserve order but allow content changes
        for i, new_block in enumerate(new_blocks):
            if i < len(template_blocks):
                # Match by position in document
                matches.append((new_block, template_blocks[i]))
            else:
                # Use last template block format as fallback
                if template_blocks:
                    matches.append((new_block, template_blocks[-1]))
                else:
                    # Create default format
                    default_block = RichFormatBlock(
                        text=new_block.text,
                        block_format=BlockFormat(type=BlockType.PARAGRAPH)
                    )
                    matches.append((new_block, default_block))
        
        return matches
    
    def _apply_block_format(self, text: str, old_block: RichFormatBlock) -> RichFormatBlock:
        """Apply old block's format to new text"""
        new_block = RichFormatBlock(
            text=text,
            block_format=old_block.block_format,
            position=old_block.position,  # Preserve position
            metadata=old_block.metadata.copy()
        )
        
        # Map character formats proportionally
        if old_block.char_formats:
            old_len = len(old_block.text)
            new_len = len(text)
            
            for old_format in old_block.char_formats:
                # Calculate proportional positions
                start_ratio = old_format.start / old_len if old_len > 0 else 0
                end_ratio = old_format.end / old_len if old_len > 0 else 1
                
                new_start = int(start_ratio * new_len)
                new_end = int(end_ratio * new_len)
                
                new_format = TextFormat(
                    start=new_start,
                    end=new_end,
                    bold=old_format.bold,
                    italic=old_format.italic,
                    underline=old_format.underline,
                    font_size=old_format.font_size,
                    font_family=old_format.font_family,
                    color=old_format.color
                )
                new_block.char_formats.append(new_format)
        
        return new_block
    
    def _analyze_structure(self, template: RichFormatTemplate) -> Dict[str, int]:
        """Analyze document structure"""
        structure = defaultdict(int)
        
        for block in template.blocks:
            block_type = block.block_format.type.value
            structure[block_type] += 1
            
            if block.block_format.heading_level:
                structure[f'heading_{block.block_format.heading_level}'] += 1
        
        return dict(structure)
    
    def _classify_block(self, text: str, structure: Dict[str, int]) -> BlockType:
        """Classify what type of block this text should be"""
        text = text.strip()
        
        # Check for heading indicators
        if len(text.split()) <= 10 and text[0].isupper():
            if text.isupper() or not text[-1] in '.!?,:;':
                return BlockType.HEADING
        
        # Check for list indicators
        if re.match(r'^[-•●○▪▫■□*+]\s', text):
            return BlockType.LIST_ITEM
        if re.match(r'^\d+[\.)]\s', text):
            return BlockType.LIST_ITEM
        
        # Check for page breaks
        if "<<<PAGE_BREAK>>>" in text:
            return BlockType.PARAGRAPH  # Treat as special paragraph
        
        # Default: paragraph
        return BlockType.PARAGRAPH
    
    def _find_matching_format(self, block_type: BlockType, template: RichFormatTemplate) -> RichFormatBlock:
        """Find a matching format from template"""
        for block in template.blocks:
            if block.block_format.type == block_type:
                return block
        
        # Default: paragraph
        return RichFormatBlock(
            text="",
            block_format=BlockFormat(type=BlockType.PARAGRAPH)
        )

# ============================================================================
# MAIN PROCESSOR
# ============================================================================

class DocumentProcessor:
    """Main document processing pipeline"""
    
    def __init__(self):
        self.extractor = TextFormatExtractor()
        self.applier = FormatApplier()
    
    def process_document(self, original_text: str, enhanced_text: str, ocr_data: Optional[Dict] = None) -> RichFormatTemplate:
        """
        Process document: extract format from original, apply to enhanced text
        
        Args:
            original_text: Original document text
            enhanced_text: AI-enhanced text
            ocr_data: Optional OCR data with layout information
        """
        # Extract format from original text (with OCR data if available)
        template = self.extractor.extract_from_text_patterns(original_text, ocr_data)
        
        # Apply format to enhanced text using semantic matching
        formatted_result = self.applier.apply_format_semantic(enhanced_text, template)
        
        return formatted_result

# ============================================================================
# DOCX EXPORTER
# ============================================================================

class WordExporter:
    """Exports RichFormatTemplate to Microsoft Word (.docx) with full formatting"""

    def export(self, template: RichFormatTemplate, output_path: str) -> str:
        """
        Export formatted template to DOCX file
        
        Args:
            template: RichFormatTemplate to export
            output_path: Path where to save the DOCX file
            
        Returns:
            Path to the exported file
        """
        doc = Document()

        for block in template.blocks:
            text = block.text.strip()
            
            # Skip completely empty blocks
            if not text:
                doc.add_paragraph("")
                continue

            block_format = block.block_format
            para = None

            # Handle different block types
            if block_format.type == BlockType.HEADING:
                # Add heading with appropriate level
                level = block_format.heading_level or 1
                level = max(1, min(level, 9))  # python-docx supports heading 1-9
                para = doc.add_heading(text, level=level)

            elif block_format.type == BlockType.LIST_ITEM:
                # Add list item
                para = doc.add_paragraph(text, style="List Bullet")
                
                # Apply indentation for nested lists
                if block_format.indent_level > 0:
                    para.paragraph_format.left_indent = Pt(block_format.indent_level * 20)

            else:
                # Add regular paragraph
                para = doc.add_paragraph(text)

            # Apply paragraph alignment
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if block_format.alignment in align_map:
                para.alignment = align_map[block_format.alignment]

            # Apply margins and spacing
            if block_format.margin_top > 0:
                para.paragraph_format.space_before = Pt(block_format.margin_top)
            if block_format.margin_bottom > 0:
                para.paragraph_format.space_after = Pt(block_format.margin_bottom)
            if block_format.margin_left > 0:
                para.paragraph_format.left_indent = Pt(block_format.margin_left)

            # Apply line height
            if block_format.line_height != 1.0:
                para.paragraph_format.line_spacing = block_format.line_height

            # Apply character formatting
            if block.char_formats and para.runs:
                # Apply formatting to the first run (simplified approach)
                run = para.runs[0]
                
                # Use the first character format if available
                if block.char_formats:
                    fmt = block.char_formats[0]
                    
                    run.bold = fmt.bold
                    run.italic = fmt.italic
                    run.underline = fmt.underline
                    run.font.name = fmt.font_family
                    run.font.size = Pt(fmt.font_size)
                    
                    # Apply color if valid hex color
                    if fmt.color.startswith("#") and len(fmt.color) == 7:
                        try:
                            r = int(fmt.color[1:3], 16)
                            g = int(fmt.color[3:5], 16)
                            b = int(fmt.color[5:7], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                        except ValueError:
                            pass  # Skip invalid color

        # Save the document
        doc.save(output_path)
        print(f"[WordExporter] Document saved to: {output_path}")
        return output_path