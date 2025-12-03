import io
import re
import base64
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2

class DocumentConverter:
    """Converter for various document formats"""
    
    def extract_text(self, file_content: bytes, file_ext: str) -> str:
        """
        Extract text from various document formats
        
        Args:
            file_content: Raw file bytes
            file_ext: File extension (.docx, .pdf, .txt)
            
        Returns:
            Extracted text content
        """
        if file_ext == '.docx' or file_ext == '.doc':
            return self._extract_from_docx(file_content)
        elif file_ext == '.pdf':
            return self._extract_from_pdf(file_content)
        elif file_ext == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def create_document(
        self, 
        content: str, 
        original_format: str = '.docx',
        output_format: str = '.docx',
        include_latex: bool = False
    ) -> bytes:
        """
        Create a document from enhanced content
        
        Args:
            content: Enhanced content (possibly with LaTeX)
            original_format: Original file format
            output_format: Desired output format
            include_latex: Whether content includes LaTeX
            
        Returns:
            Document file as bytes
        """
        if output_format == '.docx':
            return self._create_docx(content, include_latex)
        elif output_format == '.pdf':
            # For PDF, first create DOCX then convert
            # In production, you'd use pandoc or similar
            docx_bytes = self._create_docx(content, include_latex)
            # For now, return DOCX (PDF conversion requires additional tools)
            return docx_bytes
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _create_docx(self, content: str, include_latex: bool = False) -> bytes:
        """
        Create DOCX document from content
        
        Args:
            content: Enhanced content
            include_latex: Whether to preserve LaTeX formatting
            
        Returns:
            DOCX file as bytes
        """
        doc = Document()
        
        # Set document styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Process content line by line
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            # Detect headings (lines that are all caps or start with #)
            if line.isupper() and len(line.split()) <= 10:
                # Likely a heading
                heading = doc.add_heading(line, level=1)
            elif line.startswith('# '):
                # Markdown-style heading
                heading_text = line.replace('#', '').strip()
                heading_level = min(len(line) - len(line.lstrip('#')), 3)
                doc.add_heading(heading_text, level=heading_level)
            elif include_latex and ('$' in line):
                # Handle LaTeX equations
                self._add_latex_paragraph(doc, line)
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                
                # Check if it's a bullet point
                if line.startswith('- ') or line.startswith('• '):
                    para.style = 'List Bullet'
                    para.text = line[2:].strip()
                elif re.match(r'^\d+\.', line):
                    para.style = 'List Number'
                    para.text = re.sub(r'^\d+\.\s*', '', line)
        
        # Save to bytes
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        return output_buffer.getvalue()
    
    def _add_latex_paragraph(self, doc: Document, line: str):
        """
        Add paragraph with LaTeX equations
        
        For display equations ($$...$$), center them
        For inline equations ($...$), keep them inline with special formatting
        """
        # Check if it's a display equation
        if '$$' in line:
            # Display equation - center it
            equation_match = re.search(r'\$\$(.*?)\$\$', line)
            if equation_match:
                equation_text = equation_match.group(1).strip()
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(equation_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(12)
                run.italic = True
        else:
            # Inline equation or mixed text
            para = doc.add_paragraph()
            
            # Split by $ to find equations
            parts = line.split('$')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Regular text
                    if part:
                        para.add_run(part)
                else:
                    # Equation
                    run = para.add_run(part)
                    run.font.name = 'Cambria Math'
                    run.italic = True
    
    def preserve_formatting(self, original_doc: Document, enhanced_content: str) -> Document:
        """
        Attempt to preserve original document formatting
        
        Args:
            original_doc: Original document
            enhanced_content: Enhanced text content
            
        Returns:
            New document with enhanced content and preserved formatting
        """
        # This is a simplified version
        # In production, you'd want more sophisticated formatting preservation
        new_doc = Document()
        
        # Copy styles from original
        for style in original_doc.styles:
            try:
                if style.name not in new_doc.styles:
                    new_doc.styles.add_style(style.name, style.type)
            except:
                pass
        
        # Add enhanced content
        for line in enhanced_content.split('\n'):
            if line.strip():
                new_doc.add_paragraph(line)
        
        return new_doc

    def add_signature(self, file_content: bytes, signature_data: str, position: str = 'bottom-right', signer_name: str = None) -> bytes:
        """
        Add signature to document
        
        Args:
            file_content: Original DOCX content
            signature_data: Base64 encoded signature image
            position: Position of signature (bottom-right, bottom-center, bottom-left)
            signer_name: Optional name of signer
            
        Returns:
            Signed document bytes
        """
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Decode signature image
            if ',' in signature_data:
                signature_data = signature_data.split(',')[1]
            image_data = base64.b64decode(signature_data)
            image_stream = io.BytesIO(image_data)
            
            # Add some spacing
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Add signature paragraph
            sig_para = doc.add_paragraph()
            
            # Set alignment based on position
            if 'right' in position:
                sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'center' in position:
                sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # Add image
            run = sig_para.add_run()
            run.add_picture(image_stream, width=Inches(2.0))
            
            # Add signer name if provided
            if signer_name:
                name_para = doc.add_paragraph(signer_name)
                if 'right' in position:
                    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'center' in position:
                    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Save
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            return output_buffer.getvalue()
            
        except Exception as e:
            raise ValueError(f"Failed to add signature: {str(e)}")
