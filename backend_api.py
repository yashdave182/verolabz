"""
Document Tweaker Backend API
Flask backend for handling document upload, OCR extraction, AI enhancement, and format preservation
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import json
import time
import uuid
from pathlib import Path
from datetime import datetime
import requests
from io import BytesIO
import traceback

# Import Gemini SDK
import google.generativeai as genai

# Import format processor
from module import DocumentProcessor, RichFormatTemplate, TextFormatExtractor, FormatApplier

# Import docx for Word document generation
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
# Enable CORS for React frontend - configure origins from environment variable
frontend_url = os.getenv("FRONTEND_URL", "*")
# For development, also allow localhost:8080
origins = [frontend_url]
if frontend_url != "*" and frontend_url != "http://localhost:8080":
    origins.append("http://localhost:8080")

CORS(app, origins=origins)

# Configuration
UPLOAD_FOLDER = Path("uploads")
PROCESSED_FOLDER = Path("processed")
UPLOAD_FOLDER.mkdir(exist_ok=True)
PROCESSED_FOLDER.mkdir(exist_ok=True)

# API Keys from environment variables
UNSTRACT_API_KEY = os.getenv("UNSTRACT_API_KEY", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Configure Gemini
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# Unstract API Configuration
UNSTRACT_API_URL = "https://llmwhisperer-api.us-central.unstract.com/api/v2"

# For PythonAnywhere compatibility, we need to ensure the app can be imported
# without immediately starting the server
def create_app():
    """Factory function to create the Flask app"""
    return app

# ============================================================================
# UNSTRACT OCR SERVICE
# ============================================================================


class UnstractOCRService:
    """Service for extracting text from documents using Unstract LLMWhisperer API"""

    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = UNSTRACT_API_URL

    def extract_text(
        self, file_path: str, mode: str = "form", output_mode: str = "layout_preserving"
    ) -> dict:
        """
        Extract text from document using Unstract API

        Args:
            file_path: Path to the document file
            mode: Processing mode (native_text, low_cost, high_quality, form, table)
            output_mode: Output mode (layout_preserving, text)

        Returns:
            dict with status, whisper_hash, extracted text, and structured data
        """
        try:
            print(f"[Unstract] Starting extraction for: {file_path}")
            print(f"[Unstract] Mode: {mode}, Output Mode: {output_mode}")

            # Step 1: Submit document for processing
            with open(file_path, "rb") as f:
                file_data = f.read()

            headers = {"unstract-key": self.api_key}

            params = {
                "mode": mode,
                "output_mode": output_mode,
                "page_seperator": "<<<PAGE_BREAK>>>",
                "mark_vertical_lines": "true",
                "mark_horizontal_lines": "true",
                "add_line_nos": "true",
            }

            # Submit for processing
            submit_url = f"{self.base_url}/whisper"
            response = requests.post(
                submit_url, headers=headers, params=params, data=file_data, timeout=60
            )

            if response.status_code != 202:
                error_msg = (
                    f"Unstract API error: {response.status_code} - {response.text}"
                )
                print(f"[Unstract] Error: {error_msg}")
                return {"success": False, "error": error_msg}

            result = response.json()
            whisper_hash = result.get("whisper_hash")

            if not whisper_hash:
                return {"success": False, "error": "No whisper_hash received from API"}

            print(f"[Unstract] Document submitted. Hash: {whisper_hash}")

            # Step 2: Poll for status
            max_attempts = 60  # 60 attempts with 2 second intervals = 2 minutes max
            attempt = 0

            while attempt < max_attempts:
                time.sleep(2)
                attempt += 1

                status_url = f"{self.base_url}/whisper-status"
                status_params = {"whisper_hash": whisper_hash}

                status_response = requests.get(
                    status_url, headers=headers, params=status_params, timeout=30
                )

                if status_response.status_code == 200:
                    status_data = status_response.json()
                    status = status_data.get("status")

                    print(f"[Unstract] Attempt {attempt}: Status = {status}")

                    if status == "processed":
                        print(f"[Unstract] Processing complete!")
                        break
                    elif status == "processing":
                        continue
                    elif status == "failed":
                        return {"success": False, "error": "OCR processing failed"}

            if attempt >= max_attempts:
                return {"success": False, "error": "OCR processing timeout"}

            # Step 3: Retrieve extracted text
            retrieve_url = f"{self.base_url}/whisper-retrieve"
            retrieve_params = {"whisper_hash": whisper_hash}

            retrieve_response = requests.get(
                retrieve_url, headers=headers, params=retrieve_params, timeout=30
            )

            if retrieve_response.status_code != 200:
                return {
                    "success": False,
                    "error": f"Failed to retrieve text: {retrieve_response.status_code}",
                }

            extracted_text = retrieve_response.text

            # Step 4: Retrieve metadata and layout information
            layout_data = {}
            metadata = {}
            
            try:
                # Try to get layout information
                layout_url = f"{self.base_url}/whisper-layout"
                layout_params = {"whisper_hash": whisper_hash}
                
                layout_response = requests.get(
                    layout_url, headers=headers, params=layout_params, timeout=30
                )
                
                if layout_response.status_code == 200:
                    layout_data = layout_response.json()
                    print(f"[Unstract] Retrieved layout data")
            except Exception as e:
                print(f"[Unstract] Warning: Could not retrieve layout data: {str(e)}")
            
            try:
                # Try to get metadata
                metadata_url = f"{self.base_url}/whisper-metadata"
                metadata_params = {"whisper_hash": whisper_hash}
                
                metadata_response = requests.get(
                    metadata_url, headers=headers, params=metadata_params, timeout=30
                )
                
                if metadata_response.status_code == 200:
                    metadata = metadata_response.json()
                    print(f"[Unstract] Retrieved metadata")
            except Exception as e:
                print(f"[Unstract] Warning: Could not retrieve metadata: {str(e)}")

            print(f"[Unstract] Extraction complete. Text length: {len(extracted_text)}")

            return {
                "success": True,
                "text": extracted_text,
                "whisper_hash": whisper_hash,
                "layout_data": layout_data,
                "metadata": metadata
            }

        except requests.exceptions.Timeout:
            return {"success": False, "error": "API request timeout"}
        except Exception as e:
            print(f"[Unstract] Exception: {str(e)}")
            traceback.print_exc()
            return {"success": False, "error": f"Extraction error: {str(e)}"}


# ============================================================================
# GEMINI AI SERVICE
# ============================================================================


class GeminiEnhancerService:
    """Service for enhancing documents using Google Gemini AI"""

    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        # Use the Gemini 2.5 Pro model as it's the most advanced
        self.model = genai.GenerativeModel("gemini-2.5-flash")

    def enhance_document(
        self, text: str, context: str, preserve_format: bool = True
    ) -> dict:
        """
        Enhance document text using Gemini AI

        Args:
            text: Original document text
            context: User's enhancement context/instructions
            preserve_format: Whether to preserve formatting markers

        Returns:
            dict with success status and enhanced text
        """
        try:
            print(f"[Gemini] Starting enhancement. Text length: {len(text)}")
            print(f"[Gemini] Context: {context}")

            # Build prompt with improved structure and clarity
            if preserve_format:
                system_prompt = """[ROLE]
You are an expert document enhancer specializing in improving content while preserving exact formatting and layout. You are using Gemini 2.5 Pro, our most advanced thinking model.

[PRIMARY INSTRUCTION]
The user has provided a specific request. This MUST be the primary focus of your enhancement. All other improvements should support this main request.

[FORMAT PRESERVATION RULES]
1. NEVER modify formatting markers: **bold**, *italic*, <<<PAGE_BREAK>>>, [1], [2], etc.
2. NEVER change document structure: headings, paragraphs, lists, tables
3. NEVER reorganize content order
4. NEVER add or remove lines
5. Keep all special characters and symbols in their original positions
6. Maintain exact indentation and spacing patterns

[ENHANCEMENT FOCUS]
- Fix grammar and spelling errors
- Improve word choice and clarity
- Adjust tone based on user request
- Make content more professional or engaging as requested
- Follow the user's specific instructions precisely

[ADVANCED CAPABILITIES]
As Gemini 2.5 Pro, you can:
- Reason over complex problems in code, math, and STEM
- Analyze large datasets and documents using long context
- Handle sophisticated document structures

[IMPORTANT]
The document may contain special layout markers such as:
- Page breaks: <<<PAGE_BREAK>>>
- Line numbers: [1], [2], [3], etc.
- Table structures with | and - characters
- Indentation with spaces or tabs
- Special formatting markers

DO NOT change any of these markers or the document structure."""
            else:
                system_prompt = """[ROLE]
You are an expert document enhancer using Gemini 2.5 Pro. Improve the document's content based on the user's instructions while maintaining a professional and clear style."""

            # Create a more structured prompt with clear sections
            # Make the user context more prominent and actionable
            user_prompt = f"""{system_prompt}

[USER'S SPECIFIC REQUEST]
{context}

[DOCUMENT TO ENHANCE]
{text}

[CRITICAL INSTRUCTIONS]
1. The USER'S SPECIFIC REQUEST above is the PRIMARY focus - make this happen
2. Preserve ALL formatting exactly as shown
3. Return ONLY the enhanced document
4. Do not include any additional commentary or explanations"""

            # Adjust generation parameters to take advantage of Gemini 2.5 Pro capabilities
            # Default parameters optimized for Gemini 2.5 Pro
            temperature = 0.3
            top_p = 0.95
            top_k = 64
            max_output_tokens = 16384  # Take advantage of larger context window
            
            # Adjust parameters based on context keywords
            context_lower = context.lower()
            if any(word in context_lower for word in ["creative", "imaginative", "story", "narrative"]):
                temperature = 0.7  # More creative
            elif any(word in context_lower for word in ["formal", "professional", "business"]):
                temperature = 0.2  # More precise
            elif any(word in context_lower for word in ["grammar", "spelling", "fix"]):
                temperature = 0.1  # Most consistent
            elif any(word in context_lower for word in ["technical", "code", "math", "stem"]):
                temperature = 0.4  # Balanced for technical content

            # Use proper system instruction and generation configuration
            response = self.model.generate_content(
                user_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature,
                    top_p=top_p,
                    top_k=top_k,
                    max_output_tokens=max_output_tokens,
                    response_mime_type="text/plain"
                ),
            )

            if not response or not response.text:
                return {"success": False, "error": "No response from Gemini API"}

            enhanced_text = response.text.strip()

            print(f"[Gemini] Enhancement complete. Output length: {len(enhanced_text)}")
            print(f"[Gemini] Used parameters: temp={temperature}, top_p={top_p}, top_k={top_k}, max_tokens={max_output_tokens}")

            return {"success": True, "text": enhanced_text}

        except Exception as e:
            print(f"[Gemini] Exception: {str(e)}")
            traceback.print_exc()
            return {"success": False, "error": f"Enhancement error: {str(e)}"}


# ============================================================================
# API ENDPOINTS
# ============================================================================


@app.route("/api/health", methods=["GET"])
def health_check():
    """Health check endpoint"""
    return jsonify(
        {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "unstract_configured": bool(UNSTRACT_API_KEY),
            "gemini_configured": bool(GEMINI_API_KEY),
        }
    )


@app.route("/api/upload", methods=["POST"])
def upload_document():
    """Upload document and extract text using Unstract OCR"""
    try:
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400

        file = request.files["file"]

        if file.filename == "":
            return jsonify({"success": False, "error": "No file selected"}), 400

        # Validate file type
        allowed_extensions = {".pdf", ".doc", ".docx", ".txt"}
        file_ext = Path(file.filename).suffix.lower()

        if file_ext not in allowed_extensions:
            return jsonify(
                {
                    "success": False,
                    "error": f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}",
                }
            ), 400

        # Generate unique filename
        document_id = str(uuid.uuid4())
        filename = f"{document_id}{file_ext}"
        file_path = UPLOAD_FOLDER / filename

        # Save file
        file.save(file_path)
        print(f"[Upload] File saved: {file_path}")

        # Get processing mode from request
        mode = request.form.get(
            "mode", "form"
        )  # form mode for documents with structure
        output_mode = request.form.get("output_mode", "layout_preserving")

        # Handle TXT files directly
        if file_ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()

            return jsonify(
                {
                    "success": True,
                    "document_id": document_id,
                    "text": extracted_text,
                    "filename": file.filename,
                    "method": "direct",
                }
            )

        # Use Unstract OCR for PDF and Word documents
        if not UNSTRACT_API_KEY:
            return jsonify(
                {"success": False, "error": "Unstract API key not configured"}
            ), 500

        ocr_service = UnstractOCRService(UNSTRACT_API_KEY)
        result = ocr_service.extract_text(
            str(file_path), mode=mode, output_mode=output_mode
        )

        if not result["success"]:
            return jsonify(result), 500

        return jsonify(
            {
                "success": True,
                "document_id": document_id,
                "text": result["text"],
                "filename": file.filename,
                "whisper_hash": result.get("whisper_hash"),
                "method": "unstract_ocr",
            }
        )

    except Exception as e:
        print(f"[Upload] Exception: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Upload error: {str(e)}"}), 500


@app.route("/api/enhance", methods=["POST"])
def enhance_document():
    """Enhance document using Gemini AI with format preservation"""
    start_time = time.time()
    try:
        data = request.json

        if not data:
            return jsonify({"success": False, "error": "No data provided"}), 400

        original_text = data.get("text", "")
        context = data.get("context", "")
        preserve_format = data.get("preserve_format", True)

        if not original_text or not context:
            return jsonify(
                {"success": False, "error": "Both text and context are required"}
            ), 400

        # Check Gemini API key
        if not GEMINI_API_KEY:
            return jsonify(
                {"success": False, "error": "Gemini API key not configured"}
            ), 500

        # Extract format template from original text
        print("[Format] Extracting format template...")
        format_start = time.time()
        processor = DocumentProcessor()
        original_template = processor.extractor.extract_from_text_patterns(
            original_text
        )
        format_time = time.time() - format_start
        print(f"[Format] Template extraction completed in {format_time:.2f} seconds")

        # Save original template
        template_id = str(uuid.uuid4())
        template_path = PROCESSED_FOLDER / f"{template_id}_template.json"
        original_template.save(str(template_path))

        # Enhance with Gemini
        print("[Enhancement] Starting Gemini enhancement...")
        gemini_start = time.time()
        gemini_service = GeminiEnhancerService(GEMINI_API_KEY)
        enhancement_result = gemini_service.enhance_document(
            original_text, context, preserve_format=preserve_format
        )
        gemini_time = time.time() - gemini_start
        print(f"[Enhancement] Gemini processing completed in {gemini_time:.2f} seconds")

        if not enhancement_result["success"]:
            return jsonify(enhancement_result), 500

        enhanced_text = enhancement_result["text"]

        # Apply format template to enhanced text
        print("[Format] Applying format template to enhanced text...")
        apply_start = time.time()
        formatted_result = processor.applier.apply_format_semantic(
            enhanced_text, original_template
        )
        apply_time = time.time() - apply_start
        print(f"[Format] Format application completed in {apply_time:.2f} seconds")

        # Save formatted result
        result_path = PROCESSED_FOLDER / f"{template_id}_result.json"
        formatted_result.save(str(result_path))

        return jsonify(
            {
                "success": True,
                "enhanced_text": enhanced_text,
                "template_id": template_id,
                "format_preserved": preserve_format,
                "processing_time": {
                    "format_extraction": format_time,
                    "gemini_enhancement": gemini_time,
                    "format_application": apply_time,
                    "total": time.time() - start_time
                }
            }
        )

    except Exception as e:
        print(f"[Enhance] Exception: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Enhancement error: {str(e)}"}), 500


@app.route("/api/process", methods=["POST"])
def process_document():
    """Complete pipeline: Upload -> OCR -> Enhance -> Format"""
    try:
        # Step 1: Upload and OCR
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400

        file = request.files["file"]
        context = request.form.get("context", "")
        mode = request.form.get("mode", "form")
        preserve_format = request.form.get("preserve_format", "true").lower() == "true"

        if not context:
            return jsonify({"success": False, "error": "Context is required"}), 400

        print(f"[Process] Starting full pipeline for: {file.filename}")

        # Save file temporarily
        document_id = str(uuid.uuid4())
        file_ext = Path(file.filename).suffix.lower()
        filename = f"{document_id}{file_ext}"
        file_path = UPLOAD_FOLDER / filename
        file.save(file_path)

        # Extract text and layout data
        ocr_data = None
        if file_ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                original_text = f.read()
        else:
            if not UNSTRACT_API_KEY:
                return jsonify(
                    {"success": False, "error": "Unstract API key not configured"}
                ), 500

            ocr_service = UnstractOCRService(UNSTRACT_API_KEY)
            ocr_result = ocr_service.extract_text(str(file_path), mode=mode)

            if not ocr_result["success"]:
                return jsonify(ocr_result), 500

            original_text = ocr_result["text"]
            # Extract layout and metadata data for enhanced format preservation
            ocr_data = {
                "layout": ocr_result.get("layout_data", {}),
                "metadata": ocr_result.get("metadata", {})
            }

        print(f"[Process] Text extracted. Length: {len(original_text)}")

        # Extract format template with rich layout information
        processor = DocumentProcessor()
        original_template = processor.extractor.extract_from_text_patterns(
            original_text, ocr_data
        )

        # Enhance with Gemini
        if not GEMINI_API_KEY:
            return jsonify(
                {"success": False, "error": "Gemini API key not configured"}
            ), 500

        gemini_service = GeminiEnhancerService(GEMINI_API_KEY)
        enhancement_result = gemini_service.enhance_document(
            original_text, context, preserve_format=preserve_format
        )

        if not enhancement_result["success"]:
            return jsonify(enhancement_result), 500

        enhanced_text = enhancement_result["text"]

        # Apply format with enhanced layout preservation
        formatted_result = processor.applier.apply_format_semantic(
            enhanced_text, original_template
        )

        # Save results
        template_path = PROCESSED_FOLDER / f"{document_id}_template.json"
        result_path = PROCESSED_FOLDER / f"{document_id}_result.json"

        original_template.save(str(template_path))
        formatted_result.save(str(result_path))

        return jsonify(
            {
                "success": True,
                "document_id": document_id,
                "original_text": original_text,
                "enhanced_text": enhanced_text,
                "filename": file.filename,
            }
        )

    except Exception as e:
        print(f"[Process] Exception: {str(e)}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Processing error: {str(e)}"}), 500


@app.route("/api/download/<document_id>", methods=["GET"])
def download_document(document_id):
    """Download enhanced document"""
    try:
        result_path = PROCESSED_FOLDER / f"{document_id}_result.json"

        if not result_path.exists():
            return jsonify({"success": False, "error": "Document not found"}), 404

        # Check if user wants DOCX format
        format_type = request.args.get('format', 'txt')  # Default to txt for backward compatibility

        # Load format template
        template = RichFormatTemplate.load(str(result_path))

        if format_type.lower() == 'docx':
            # Generate DOCX file
            doc = Document()
            
            # Add content to document based on format template
            for block in template.blocks:
                # Determine block type and add appropriate content
                if block.block_format.type.name == 'HEADING' and block.block_format.heading_level:
                    # Add heading
                    heading_level = min(block.block_format.heading_level, 6)  # python-docx supports heading 1-6
                    heading = doc.add_heading(block.text, level=heading_level)
                    
                    # Apply character formatting if available
                    if block.char_formats:
                        for char_format in block.char_formats:
                            # Basic formatting support
                            if char_format.bold:
                                heading.runs[0].bold = True
                            if char_format.italic:
                                heading.runs[0].italic = True
                else:
                    # Add paragraph
                    paragraph = doc.add_paragraph(block.text)
                    
                    # Apply paragraph formatting
                    if block.block_format.alignment == 'center':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif block.block_format.alignment == 'right':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif block.block_format.alignment == 'justify':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Apply character formatting if available
                    if block.char_formats:
                        for char_format in block.char_formats:
                            # Basic character formatting support
                            if char_format.bold:
                                paragraph.runs[0].bold = True
                            if char_format.italic:
                                paragraph.runs[0].italic = True

            # Save to BytesIO buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return send_file(
                buffer,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"enhanced_document_{document_id}.docx",
            )
        else:
            # Original TXT format
            # Convert to plain text
            text_output = []
            for block in template.blocks:
                text_output.append(block.text)

            full_text = "\n\n".join(text_output)

            # Create file buffer
            buffer = BytesIO()
            buffer.write(full_text.encode("utf-8"))
            buffer.seek(0)

            return send_file(
                buffer,
                mimetype="text/plain",
                as_attachment=True,
                download_name=f"enhanced_document_{document_id}.txt",
            )

    except Exception as e:
        print(f"[Download] Exception: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    print("=" * 80)
    print("Document Tweaker Backend API")
    print("=" * 80)
    print(f"Unstract API Key: {'✓ Configured' if UNSTRACT_API_KEY else '✗ Missing'}")
    print(f"Gemini API Key: {'✓ Configured' if GEMINI_API_KEY else '✗ Missing'}")
    print("=" * 80)

    # Get port from environment variable or default to 5000
    port = int(os.environ.get("PORT", 5000))
    
    # Development server
    app.run(host="0.0.0.0", port=port, debug=False)
